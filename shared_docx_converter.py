#!/usr/bin/env python3
"""
Shared DOCX conversion utility for web scraping feasibility analysis reports.
Combines the best features from existing conversion scripts.
"""

import os
import re
import sys
from pathlib import Path
from typing import Optional, List

# Try to import docx library first (preferred), fallback to pypandoc
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    USE_DOCX = True
except ImportError:
    USE_DOCX = False

try:
    import pypandoc
    USE_PYPANDOC = True
except ImportError:
    USE_PYPANDOC = False

def markdown_to_docx_native(md_file_path: str, docx_file_path: str) -> None:
    """Convert markdown to DOCX using python-docx library (preferred method)."""
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    doc = Document()
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Handle different markdown elements
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(line[3:], style='List Number')
            
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
            
        elif line.startswith('| '):
            doc.add_paragraph(line)
            
        else:
            # Regular paragraph with inline formatting
            paragraph_text = line
            
            # Clean up markdown formatting for Word
            paragraph_text = re.sub(r'\*\*(.*?)\*\*', r'\1', paragraph_text)
            paragraph_text = re.sub(r'`(.*?)`', r'\1', paragraph_text)
            paragraph_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', paragraph_text)
            
            if paragraph_text:
                doc.add_paragraph(paragraph_text)
    
    doc.save(docx_file_path)

def markdown_to_docx_pypandoc(md_file_path: str, docx_file_path: str) -> None:
    """Convert markdown to DOCX using pypandoc (fallback method)."""
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pypandoc.convert_text(
        content,
        'docx',
        format='md',
        outputfile=docx_file_path,
        extra_args=['--standalone']
    )

def convert_md_to_docx(md_file_path: str, output_dir: Optional[str] = None) -> str:
    """
    Convert a Markdown file to DOCX format.
    
    Args:
        md_file_path: Path to the markdown file
        output_dir: Optional output directory (defaults to same as input)
        
    Returns:
        Path to the generated DOCX file
        
    Raises:
        FileNotFoundError: If the markdown file doesn't exist
        RuntimeError: If no conversion library is available
    """
    
    if not os.path.exists(md_file_path):
        raise FileNotFoundError(f"Markdown file not found: {md_file_path}")
    
    if not USE_DOCX and not USE_PYPANDOC:
        raise RuntimeError("No DOCX conversion library available. Install python-docx or pypandoc.")
    
    if output_dir is None:
        output_dir = os.path.dirname(md_file_path)
    
    base_name = os.path.splitext(os.path.basename(md_file_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}.docx")
    
    print(f"Converting '{md_file_path}' → '{output_path}' ...")
    
    try:
        if USE_DOCX:
            markdown_to_docx_native(md_file_path, output_path)
        else:
            markdown_to_docx_pypandoc(md_file_path, output_path)
        
        print("✅ Conversion complete.")
        return output_path
        
    except Exception as e:
        print(f"❌ Conversion failed: {str(e)}")
        raise

def convert_all_md_files_in_directory(directory: str) -> List[str]:
    """
    Convert all markdown files in a directory to DOCX format.
    
    Args:
        directory: Path to directory containing markdown files
        
    Returns:
        List of generated DOCX file paths
    """
    
    directory = Path(directory)
    converted_files = []
    
    for md_file in directory.glob('*.md'):
        try:
            docx_file = convert_md_to_docx(str(md_file))
            converted_files.append(docx_file)
        except Exception as e:
            print(f"Failed to convert {md_file}: {str(e)}")
    
    return converted_files

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python shared_docx_converter.py <input_file.md> [output_dir]")
        print("   or: python shared_docx_converter.py --dir <directory>")
        sys.exit(1)
    
    if sys.argv[1] == "--dir":
        if len(sys.argv) < 3:
            print("Error: Directory path required when using --dir")
            sys.exit(1)
        
        directory = sys.argv[2]
        converted = convert_all_md_files_in_directory(directory)
        print(f"Converted {len(converted)} files to DOCX format!")
        
    else:
        md_file = sys.argv[1]
        output_dir = sys.argv[2] if len(sys.argv) > 2 else None
        convert_md_to_docx(md_file, output_dir)
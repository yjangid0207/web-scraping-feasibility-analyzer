#!/usr/bin/env python3
"""
Convert Markdown files to DOCX format for better sharing and presentation.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

def markdown_to_docx(md_file_path, docx_file_path):
    """Convert a markdown file to DOCX format."""
    
    # Read the markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Process markdown line by line
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - add paragraph break
            doc.add_paragraph()
            continue
        
        # Handle different markdown elements
        if line.startswith('# '):
            # Main heading
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        elif line.startswith('## '):
            # Subheading
            doc.add_heading(line[3:], level=2)
            
        elif line.startswith('### '):
            # Sub-subheading
            doc.add_heading(line[4:], level=3)
            
        elif line.startswith('#### '):
            # Level 4 heading
            doc.add_heading(line[5:], level=4)
            
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')
            
        elif re.match(r'^\d+\. ', line):
            # Numbered list
            p = doc.add_paragraph(line[3:], style='List Number')
            
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            
        elif line.startswith('---'):
            # Horizontal rule - add a page break or line
            doc.add_paragraph('_' * 50)
            
        elif line.startswith('| '):
            # Table row - basic handling
            # For now, just add as regular text
            doc.add_paragraph(line)
            
        else:
            # Regular paragraph
            # Handle inline formatting
            paragraph_text = line
            
            # Convert **bold** text
            paragraph_text = re.sub(r'\*\*(.*?)\*\*', r'\1', paragraph_text)
            
            # Convert `code` text  
            paragraph_text = re.sub(r'`(.*?)`', r'\1', paragraph_text)
            
            # Remove markdown links but keep text
            paragraph_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', paragraph_text)
            
            if paragraph_text:
                p = doc.add_paragraph(paragraph_text)
    
    # Save the document
    doc.save(docx_file_path)
    print(f"Converted {md_file_path} to {docx_file_path}")

def convert_all_md_files(directory):
    """Convert all markdown files in a directory to DOCX."""
    directory = Path(directory)
    
    for md_file in directory.glob('*.md'):
        docx_file = directory / f"{md_file.stem}.docx"
        markdown_to_docx(md_file, docx_file)

if __name__ == "__main__":
    # Convert all MD files in current directory
    current_dir = Path(__file__).parent
    convert_all_md_files(current_dir)
    print("All markdown files have been converted to DOCX format!")